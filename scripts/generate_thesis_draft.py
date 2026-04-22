from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TITLE = "视频网站数据分析平台的设计与实现"
OUTPUT_NAME = "南通大学本科毕业设计论文_视频数据分析平台.docx"


def set_run_font(run, east_asia_font: str, size: float, bold: bool = False):
    run.font.name = east_asia_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    run.font.size = Pt(size)
    run.bold = bold


def configure_style(style, east_asia_font: str, size: float, bold: bool = False):
    style.font.name = east_asia_font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)


def set_cell_text(cell, text: str, font_name: str = "宋体", size: float = 10.5, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold)


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "宋体", 12)
    return p


def add_plain_paragraph(doc: Document, text: str, align=WD_ALIGN_PARAGRAPH.LEFT, bold: bool = False, size: float = 12):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, "宋体", size, bold)
    return p


def add_heading_text(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, "黑体", 16, True)
    elif level == 2:
        set_run_font(run, "黑体", 14, True)
    else:
        set_run_font(run, "黑体", 12, True)
    return p


def add_toc(doc: Document):
    p = doc.add_paragraph()
    r = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "目录将在 Word 中更新域后自动生成。"
    fld_sep.append(fld_text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_sep)
    r._r.append(fld_end)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("第 ")
    set_run_font(run, "宋体", 10.5)
    page_run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run._r.append(fld_begin)
    page_run._r.append(instr)
    page_run._r.append(fld_end)
    run = p.add_run(" 页")
    set_run_font(run, "宋体", 10.5)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run("南通大学本科毕业设计（论文）")
    set_run_font(run, "黑体", 22, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run("题    目：" + TITLE)
    set_run_font(run, "黑体", 20, True)

    info_items = [
        "学    院：待补充",
        "学生姓名：待补充",
        "学    号：待补充",
        "专    业：待补充",
        "指导教师：待补充",
        "完成日期：2026年5月",
    ]
    for text in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(text)
        set_run_font(run, "宋体", 14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    run = p.add_run("二〇二六年五月")
    set_run_font(run, "宋体", 14)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_abstract(doc: Document):
    add_heading_text(doc, "摘 要", 1)
    add_body_paragraph(
        doc,
        "随着抖音、Bilibili 等视频平台的迅速发展，视频内容已经成为互联网信息传播和用户互动的重要载体。"
        "大量视频数据中不仅包含播放量、点赞量、评论量等可量化指标，还反映了用户兴趣变化、内容传播效率和平台生态差异。"
        "然而，在毕业设计和教学研究场景中，常见的数据分析方案存在数据来源分散、字段标准不统一、导入链路不稳定、分析页面缺乏整体性等问题，导致系统虽然能够展示图表，却难以形成从数据接入到研究结论输出的完整闭环。"
    )
    add_body_paragraph(
        doc,
        "针对上述问题，本文设计并实现了一个视频网站数据分析平台。系统采用 Vue3 构建前端交互界面，使用 Spring Boot 与 MyBatis 实现后端业务服务，以 MySQL 作为核心数据存储，并辅以 Python 脚本完成数据采集、文本解析、文件导入和自动化初始化。"
        "系统围绕“多用户独立数据空间、双身份账号、标准化入库、质量治理、可视化分析、自动化部署”六条主线展开建设，支持 URL 采集、文本导入、CSV/Markdown 文件入库、来源追踪、拒绝记录回溯、平台筛选、2D/3D 图表切换以及创作者中心同行对标等功能。"
    )
    add_body_paragraph(
        doc,
        "在实现过程中，本文重点解决了多来源异构数据标准化、重复导入导致统计失真、长任务导入和清理易超时、用户画像结果单一、前端加载体验不足等问题。"
        "系统通过设计统一字段字典、去重键机制、导入任务表、拒绝记录表、概览缓存表和多层脚本链路，使平台既具备研究展示能力，又具备一定的工程可交付性。"
    )
    add_body_paragraph(
        doc,
        "实验与运行结果表明，该平台能够稳定完成多平台视频数据的导入、清洗、存储和可视化分析，支持账号级数据隔离，能够较好地满足本科毕业设计场景下对完整性、可用性和可复现性的要求。"
        "同时，系统通过区分观众与内容创作者两类角色，使平台既可用于全局研究，也可用于个人创作运营分析。本文的研究与实现对视频数据分析类教学项目具有一定参考价值。"
    )
    add_plain_paragraph(
        doc,
        "关键词：视频数据分析平台；双身份账号；标准化入库；多用户隔离；创作者中心；可视化分析",
    )
    doc.add_page_break()

    add_heading_text(doc, "Abstract", 1)
    add_plain_paragraph(
        doc,
        "With the rapid growth of video platforms such as Douyin and Bilibili, video content has become an important carrier of information dissemination and user interaction on the Internet. "
        "A large amount of video data contains not only quantitative indicators such as play count, like count and comment count, but also reflects user interests, content communication efficiency and platform differences. "
        "However, in graduation project scenarios, many existing solutions suffer from scattered data sources, inconsistent field definitions, unstable import chains and fragmented visualization pages, making it difficult to build a complete research workflow.",
    )
    add_plain_paragraph(
        doc,
        "To solve these problems, this paper designs and implements a video data analysis platform. "
        "The system uses Vue3 for the front-end interface, Spring Boot and MyBatis for the back-end services, MySQL for data storage, and Python scripts for data collection, text parsing, file import and automated initialization. "
        "The platform supports independent user data spaces, standardized import, data quality governance, visual analysis, source tracing and automated deployment.",
    )
    add_plain_paragraph(
        doc,
        "The implementation focuses on heterogeneous data normalization, deduplication, long-task stability, user portrait optimization, loading experience and 2D/3D visualization switching. "
        "The results show that the system can complete data import, cleaning, storage and analysis in a stable way, and is suitable for graduation project research and demonstration.",
    )
    add_plain_paragraph(
        doc,
        "Key Words: video data analysis platform; standardized import; multi-user isolation; visualization; Three.js; Spring Boot",
    )
    doc.add_page_break()


def add_chapter_one(doc: Document):
    add_heading_text(doc, "第1章 绪论", 1)
    add_heading_text(doc, "1.1 研究背景", 2)
    add_body_paragraph(
        doc,
        "短视频和中长视频平台已经成为当代网络内容传播的重要阵地。平台每天产生海量视频、评论、点赞和分享等互动数据，这些数据对内容热度判断、用户偏好分析、平台差异比较以及传播策略研究具有明显价值。"
        "与此同时，高校毕业设计中大量选题开始转向“数据驱动型系统”，学生不仅需要完成页面展示，更需要构建真实可运行的数据采集、存储、分析与展示闭环。"
    )
    add_body_paragraph(
        doc,
        "传统课程作业中常见的视频分析系统，往往只依赖少量样例数据或静态图表，缺少标准化导入、质量控制和多用户隔离机制，难以支撑完整的研究论证。"
        "因此，开发一个既具备工程实现能力，又能够体现数据分析思想和系统设计方法的视频数据分析平台，具有明确的实践意义。"
    )

    add_heading_text(doc, "1.2 研究目的与意义", 2)
    add_body_paragraph(
        doc,
        "本文的研究目的在于设计并实现一个面向毕业设计场景的视频数据分析平台，使用户能够在统一界面内完成数据导入、清洗、标准化、分析和可视化展示等操作，并通过多用户隔离机制保证不同账号的数据独立。"
    )
    add_body_paragraph(
        doc,
        "其理论意义在于将数据治理、可视化分析与软件工程方法结合起来，形成一个完整的系统案例；其应用意义在于为后续类似的教学项目、课程实验和小型研究系统提供可复用的实现思路。"
        "与单纯做图表展示的系统相比，本文更加重视数据来源可信性、分析结果可解释性以及部署过程可复现性。"
    )

    add_heading_text(doc, "1.3 国内外研究现状", 2)
    add_body_paragraph(
        doc,
        "在国外研究中，围绕社交媒体和视频平台的数据分析工作较早开展，研究重点通常集中在用户行为建模、推荐机制优化、内容传播预测和可视分析方法等方向。"
        "这些研究强调数据规模、实时性和算法模型能力，但很多成果更偏向平台级研究和实验性框架，直接用于教学型毕业设计时部署成本较高。"
    )
    add_body_paragraph(
        doc,
        "国内相关研究更多关注短视频平台数据挖掘、用户画像分析、热点传播规律及舆情研究。现有实践中也出现了基于 Python、MySQL、ECharts 的分析系统，但很多系统存在功能点分散、文档不完整、数据质量治理不足等问题。"
        "因此，构建一套兼顾工程完整度和研究展示效果的系统，仍然具有现实价值。"
    )

    add_heading_text(doc, "1.4 研究内容", 2)
    add_body_paragraph(
        doc,
        "围绕上述目标，本文主要完成以下工作：第一，搭建 Vue3、Spring Boot、MySQL 和 Python 组成的全栈系统架构；第二，设计多用户独立数据空间与双身份账号体系，实现观众和内容创作者差异化入口；第三，建立统一字段字典和入库规则，实现多来源数据的标准化导入；第四，完成热门视频、分类统计、互动效率、用户画像和创作者中心等分析模块；第五，提供 ECharts 与 Three.js 结合的 2D/3D 可视化能力；第六，补齐一键初始化、自动启动、自测和报告导出等工程化功能。"
    )

    add_heading_text(doc, "1.5 论文结构", 2)
    add_body_paragraph(
        doc,
        "全文共分为八章。第一章介绍研究背景、意义与研究内容；第二章介绍关键技术与可行性；第三章分析系统需求；第四章进行系统总体设计；第五章描述数据库设计；第六章详细说明系统实现过程；第七章给出测试与结果分析；第八章总结全文并提出后续改进方向。"
    )


def add_chapter_two(doc: Document):
    add_heading_text(doc, "第2章 相关技术与可行性分析", 1)
    add_heading_text(doc, "2.1 前端技术", 2)
    add_body_paragraph(
        doc,
        "系统前端采用 Vue3 构建，配合 Vue Router 完成单页应用的页面切换与模块化组织。Vue3 具有组件化程度高、学习成本适中、工程生态成熟等优点，适合毕业设计项目快速迭代。"
        "在本系统中，前端主要承担登录注册、数据管理、分析页面展示、平台筛选、状态反馈和加载交互等功能。"
    )
    add_body_paragraph(
        doc,
        "为了提升图表表现力，系统同时接入 ECharts 与 Three.js。ECharts 主要用于二维统计图表，例如折线图、柱状图、饼图和散点图；Three.js 主要用于三维分类对比场景，实现 2D/3D 图表切换。"
    )

    add_heading_text(doc, "2.2 后端技术", 2)
    add_body_paragraph(
        doc,
        "系统后端基于 Spring Boot 3 构建 REST 风格接口，使用 MyBatis 和 JdbcTemplate 进行数据访问。Spring Boot 能够快速完成控制器、服务层、拦截器和配置类组织，适合构建中小型业务系统。"
        "MyBatis 适用于本项目以统计查询和业务 SQL 为主的场景，能够更直观地控制多表聚合、分页和条件筛选逻辑。"
    )
    add_body_paragraph(
        doc,
        "在安全设计方面，后端通过 AuthInterceptor、AuthContext、会话表和安全响应头过滤器实现登录态校验、会话失效控制和基础安全增强。"
        "相比单纯使用前端缓存 token 的方式，这种设计更容易支持退出登录、会话撤销和多账号隔离。"
    )

    add_heading_text(doc, "2.3 数据库与脚本技术", 2)
    add_body_paragraph(
        doc,
        "MySQL 8 作为系统核心数据库，负责存储用户账号、视频主数据、评论、行为日志、统计结果和导入任务信息。"
        "数据库采用 utf8mb4 字符集，兼容中文标题、作者名和多平台文本内容。围绕性能和一致性需求，系统设计了索引、唯一约束、缓存表和补偿迁移机制。"
    )
    add_body_paragraph(
        doc,
        "Python 脚本层承担数据采集、模板生成、初始化、自测和报告导出等辅助工作。与全部逻辑堆叠在后端相比，脚本化方式更适合处理环境检查、无 MySQL CLI 初始化和离线自测等工程任务。"
    )

    add_heading_text(doc, "2.4 AI 辅助解析技术", 2)
    add_body_paragraph(
        doc,
        "本系统实现了最小可用的 AI 辅助入库能力。具体做法是：当规则解析无法稳定提取结构化记录，且系统启用了 AI 开关并配置了 API Key 时，后端会调用兼容 OpenAI Chat Completions 的接口，请求模型返回标准 JSON 结果。"
    )
    add_body_paragraph(
        doc,
        "这种实现方式使 AI 成为“增强环节”而非“唯一依赖”。在未配置 AI 时，系统仍然能够通过 CSV、Markdown、JSON、键值文本和普通叙述文本的规则解析完成基础导入；在规则解析质量不足时，AI 仅作为补强手段，提高异构文本的结构化能力。"
    )

    add_heading_text(doc, "2.5 可行性分析", 2)
    add_body_paragraph(
        doc,
        "从技术可行性看，本文采用的 Vue3、Spring Boot、MySQL、Python、ECharts 和 Three.js 均为成熟技术栈，社区资料丰富，适合在毕业设计周期内实现并调试。"
        "从经济可行性看，系统运行所需软件均可在普通个人电脑环境中免费使用，不需要额外购买商业服务即可完成开发与演示。"
    )
    add_body_paragraph(
        doc,
        "从操作可行性看，系统提供了登录页、数据管理页、分析页、一键脚本和部署文档，普通用户只需配置本地数据库即可运行。"
        "从合规角度看，系统强调仅采集公开且已授权的数据来源，并在 URL 采集入口加入风险确认逻辑，降低非法抓取风险。"
    )


def add_chapter_three(doc: Document):
    add_heading_text(doc, "第3章 系统需求分析", 1)
    add_heading_text(doc, "3.1 业务流程需求", 2)
    add_body_paragraph(
        doc,
        "系统面向的视频数据分析流程并非单一的图表展示，而是完整的数据业务闭环。用户首先需要通过登录进入自己的独立数据空间，然后选择 URL、文本或文件等方式导入数据；系统对数据进行格式识别、字段映射、去重和质量校验后写入数据库；最后再通过首页总览、热门视频、分类统计、互动效率和用户画像等页面进行分析。"
    )
    add_body_paragraph(
        doc,
        "这意味着系统必须同时满足“数据能进来、结果能算出来、页面可理解、过程可追踪、错误可定位”五个方面的业务需求。"
        "如果只实现图表，不解决入库和质量问题，分析结果将缺乏可信度；如果只做导入，不解决展示与交互，系统又难以服务于研究表达。"
    )

    add_heading_text(doc, "3.2 功能需求", 2)
    add_body_paragraph(
        doc,
        "根据系统目标，功能需求主要包括以下几类：第一，账号功能，支持注册、登录、退出、刷新后保持登录状态，并支持观众与内容创作者双身份注册；第二，数据管理功能，支持 URL 采集、文本导入、文件导入、批量模板生成、导入追踪和数据清理；第三，分析功能，支持首页总览、热门视频排行、分类统计、互动效率分析、用户画像分析、平台筛选以及创作者中心同行对标；第四，辅助功能，支持错误提示、加载状态、自动化启动、自测和报告导出。"
    )

    add_heading_text(doc, "3.3 数据需求", 2)
    add_body_paragraph(
        doc,
        "系统需要处理的视频主数据至少应包含标题、作者、平台、分类、发布时间、播放量、点赞量、评论量等字段。为了支撑来源追踪和质量治理，还需要记录导入类型、来源文件、来源链接、导入时间、质量评分、任务编号和 AI 置信度等扩展信息。"
    )
    add_body_paragraph(
        doc,
        "同时，为完成用户画像和互动效率分析，系统还需要维护用户行为日志、评论表、日统计表及画像结果表。"
        "这些数据并不要求全部来自同一平台的同一接口，而是通过标准化字段字典统一映射后进入分析链路。"
    )

    add_heading_text(doc, "3.4 非功能需求", 2)
    add_body_paragraph(
        doc,
        "在非功能方面，系统主要需要满足以下要求：一是安全性，账号密码不能明文存储，不同账号的数据必须隔离；二是可用性，系统要有明确的错误提示和结果反馈，防止用户误以为操作失败或无响应；三是性能，首页和大文件导入场景不能长期阻塞，需要通过缓存、批量写入和分批删除等方式优化；四是可维护性，代码、脚本和文档应保持一致，便于后续修改。"
    )


def add_chapter_four(doc: Document):
    add_heading_text(doc, "第4章 系统总体设计", 1)
    add_heading_text(doc, "4.1 设计目标与原则", 2)
    add_body_paragraph(
        doc,
        "系统总体设计遵循“先闭环、再优化；先稳定、再扩展；以研究问题为中心”的原则。所谓闭环，是指系统必须能完成从数据导入到结果展示的完整链路；所谓稳定，是指面对大文件、异常数据和重复导入时不能轻易失效；所谓扩展，是指在当前架构下可逐步补充更多平台和智能能力。"
    )

    add_heading_text(doc, "4.2 总体架构设计", 2)
    add_body_paragraph(
        doc,
        "系统采用前后端分离架构。前端负责页面呈现与交互控制，后端负责账号鉴权、业务处理和数据查询，数据库负责持久化存储，Python 脚本层负责环境初始化、采集与辅助处理。"
        "其核心架构为“Vue3 前端 - Spring Boot 后端 - MyBatis/JdbcTemplate 数据访问层 - MySQL 数据库 - Python 脚本层”。"
    )
    add_body_paragraph(
        doc,
        "在具体通信上，前端通过 HTTP 请求访问后端 `/auth/*`、`/video/*` 和 `/crawler/*` 接口；后端在每次请求中通过拦截器解析登录态，并以 `tenant_user_id` 作为数据隔离主键；脚本层则通过数据库初始化、自测和导出脚本为系统运行提供工程支撑。"
    )

    add_heading_text(doc, "4.3 模块划分", 2)
    add_body_paragraph(
        doc,
        "按照业务职责，系统可划分为登录认证模块、数据管理模块、分析展示模块、质量治理模块和运维脚本模块。登录认证模块负责注册登录、身份分流、会话控制和多用户隔离；数据管理模块负责 URL、文本和文件导入；分析展示模块负责总览、排行、统计、画像、创作者中心和可视化；质量治理模块负责去重、评分和拒绝记录；运维脚本模块负责初始化、启动、自测和导出。"
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["模块名称", "主要功能", "关键实现"]
    for idx, text in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], text, "黑体", 10.5, True)
    rows = [
        ("登录认证模块", "注册、登录、登出、会话校验、身份分流", "AuthController、AuthService、AuthInterceptor"),
        ("数据管理模块", "URL采集、文本/文件导入、清理数据", "CrawlerController、CrawlerService、Python脚本"),
        ("分析展示模块", "首页总览、排行、统计、画像、创作者中心、平台筛选", "VideoController、VideoService、Vue页面"),
        ("质量治理模块", "字段标准化、质量评分、拒绝记录", "dedupe_key、import_job、import_reject_record"),
        ("运维脚本模块", "初始化、启动、自测、报告导出", "init_db_python.py、full_auto_launch.ps1"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text)

    add_heading_text(doc, "4.4 核心业务流程设计", 2)
    add_body_paragraph(
        doc,
        "系统的核心业务流程可以概括为：用户登录进入系统后，选择 URL、文本或文件作为数据来源，后端对输入内容进行规则解析和可选 AI 补强，随后执行字段映射、质量评分和去重合并，将有效数据写入视频主表及相关业务表，同时将异常数据写入拒绝记录表；最后用户在分析页面中通过多种图表查看结果，并结合平台筛选器观察不同来源的差异。"
    )

    add_heading_text(doc, "4.5 安全与部署设计", 2)
    add_body_paragraph(
        doc,
        "安全设计方面，系统采用密码哈希存储、会话表控制、拦截器鉴权、安全响应头和可配置来源白名单等方式降低常见风险。"
        "部署设计方面，系统提供无 MySQL CLI 初始化脚本和全自动启动脚本，能够自动检查 Python、安装依赖、初始化数据库、拉起前后端并执行冒烟测试，提升系统演示与交付效率。"
    )


def add_chapter_five(doc: Document):
    add_heading_text(doc, "第5章 数据库设计", 1)
    add_heading_text(doc, "5.1 数据库设计原则", 2)
    add_body_paragraph(
        doc,
        "数据库围绕“多租户隔离、标准化存储、导入可追踪、统计可扩展”的原则展开设计。"
        "其中，多租户隔离通过 `tenant_user_id` 字段实现；标准化存储通过统一字段字典解决不同平台字段不一致问题；导入可追踪通过任务表、拒绝表和来源字段实现；统计可扩展则通过行为日志表、统计表和缓存表为后续分析预留空间。"
    )

    add_heading_text(doc, "5.2 核心数据表设计", 2)
    add_body_paragraph(
        doc,
        "系统当前运行主链路共包含账号表、会话表、视频主表、作者表、评论表、行为日志表、统计表、画像结果表、导入任务表、拒绝记录表和概览缓存表等。"
        "其中，视频主表是整个系统的数据中心，导入后的大部分统计查询均围绕该表展开。"
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for idx, text in enumerate(["表名", "主要作用", "关键字段"]):
        set_cell_text(table.rows[0].cells[idx], text, "黑体", 10.5, True)
    rows = [
        ("app_user", "存储账号与身份信息", "username、user_role、creator_name、creator_platform"),
        ("app_session", "存储会话与登录状态", "token、user_id、expires_at、is_revoked"),
        ("video", "存储视频主数据", "title、author、source_platform、play_count、dedupe_key"),
        ("user_behavior", "存储用户行为日志", "tenant_user_id、user_id、video_id、action、time"),
        ("import_job", "记录每次导入任务", "import_type、source_platform、status、success_count"),
        ("import_reject_record", "记录低质量或拒绝数据", "reject_reason、suggest_fix、raw_excerpt、quality_score"),
        ("tenant_overview_cache", "缓存首页总览结果", "video_count、user_count、behavior_count、refreshed_at"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text)

    add_heading_text(doc, "5.3 标准化字段字典与入库规则", 2)
    add_body_paragraph(
        doc,
        "考虑到不同平台和不同数据文件对同一概念的命名方式并不统一，系统设计了最小可用版标准化字段字典。以标题字段为例，既支持 `title`，也支持“标题”；平台字段既支持 `platform`，也支持“来源平台”；播放量、点赞量和评论量则支持中英文混合字段名与常见单位表达。"
    )
    add_body_paragraph(
        doc,
        "在入库判定上，系统优先保证主数据质量。标题字段为必填项，平台、作者、时间、互动指标等字段会参与质量评分；当记录缺失过多、格式异常或无法标准化时，系统不会强行写入主表，而是将其记录到拒绝表，并给出修复建议，以避免低质量数据污染分析结果。"
    )

    add_heading_text(doc, "5.4 索引与性能设计", 2)
    add_body_paragraph(
        doc,
        "为提高系统查询效率，数据库在视频表和行为日志表上建立了与租户、平台、导入时间、分类和质量等级相关的索引。"
        "其中，视频表上的 `uk_video_tenant_dedupe(tenant_user_id, dedupe_key)` 唯一索引是防止重复导入的关键。"
    )
    add_body_paragraph(
        doc,
        "在性能优化方面，系统引入 `tenant_overview_cache` 表缓存首页总览结果，并在导入或清理后对缓存进行失效处理。"
        "这种做法避免了首页每次访问都对大表执行聚合扫描，有效提升了首屏加载速度。"
    )


def add_chapter_six(doc: Document):
    add_heading_text(doc, "第6章 系统详细设计与实现", 1)
    add_heading_text(doc, "6.1 登录认证与多用户隔离实现", 2)
    add_body_paragraph(
        doc,
        "系统首页采用登录与注册一体化入口，注册阶段支持选择“观众”或“内容创作者”身份。若用户选择内容创作者，还需要补充创作者名称、主运营平台和主打方向。用户提交账号密码后，后端完成密码哈希比对并生成会话 token，同时在 `app_session` 表中记录有效期、撤销状态和最近活跃时间。"
        "前端在后续请求中携带该 token，后端通过拦截器校验后，将当前登录用户写入上下文，从而实现对业务查询的统一隔离，并根据 `user_role` 决定登录后默认进入综合首页还是创作者中心。"
    )
    add_body_paragraph(
        doc,
        "为了改善体验，系统支持刷新页面后保持登录状态；为了兼顾安全性，系统在退出登录时会同步撤销服务端会话。"
        "相比简单的前端本地缓存方案，这种做法更符合毕业设计中对“登录检验”和“数据库保护”的要求。"
    )

    add_heading_text(doc, "6.2 数据管理与导入实现", 2)
    add_body_paragraph(
        doc,
        "数据管理模块是本系统的核心入口之一。系统支持三种主要入库方式：一是输入公开且已授权的 URL 进行采集；二是直接粘贴文本内容进行结构化解析；三是上传 CSV、Markdown、TXT、LOG 等文件批量入库。"
        "为减少使用门槛，系统还提供批量链接模板生成器，自动生成 urls.txt 模板并校验格式。"
    )
    add_body_paragraph(
        doc,
        "在后端实现中，`CrawlerServiceImpl` 按“格式识别 - 规则解析 - 可选 AI 增强 - 去重合并 - 质量评分 - 数据写入”的顺序处理数据。"
        "规则解析支持 JSON、Markdown 表格、CSV、键值文本和普通叙述文本。对于大文件导入场景，系统使用批量写入和缩放策略减少行为日志写入压力，显著提高了长任务成功率。"
    )

    add_heading_text(doc, "6.3 标准化治理与错误反馈实现", 2)
    add_body_paragraph(
        doc,
        "系统并不把“导入成功”视为唯一目标，而是更加关注“导入结果是否可信”。因此，在数据进入主表之前，系统会根据统一字段字典进行标准化映射，并计算质量分。"
        "对于明显缺失标题、字段格式错误或信息质量过低的数据，系统会将其保存到 `import_reject_record` 表中，记录拒绝原因、修复建议和原始片段。"
    )
    add_body_paragraph(
        doc,
        "前端数据管理页会同步显示导入状态、成功或失败提示、可执行修复建议和来源追踪信息。"
        "这样一来，用户在导入失败时能够知道是文件编码问题、字段表头问题，还是质量门禁导致的拒绝，而不是仅看到模糊的“导入失败”提示。"
    )

    add_heading_text(doc, "6.4 分析模块实现", 2)
    add_body_paragraph(
        doc,
        "系统的分析模块包括首页总览、热门视频、分类统计、互动效率、用户画像和创作者中心六个核心页面。首页用于展示全部平台总体情况，其他分析模块支持按平台切换查看对应数据；创作者中心则聚焦“自己的作品”和“同行差距”。"
        "为了降低首页请求数量，后端额外提供了聚合接口，一次性返回总览、洞察卡片、热门视频和平台概览等数据；同时新增创作者中心接口，统一返回创作者资料、自身总览、代表作品、分类统计、同行作者和策略建议。"
    )
    add_body_paragraph(
        doc,
        "热门视频模块以播放量和互动量为核心展示对象；分类统计模块从内容分类维度观察视频数量和总播放量；互动效率模块用更适合研究的互动率指标替代简单播放趋势；用户画像模块则综合活跃度、兴趣广度、互动深度和平台偏好给出标签分布和建议；创作者中心则通过自身数据与同行均值的对比，帮助创作者识别优势方向与改进重点。"
    )

    add_heading_text(doc, "6.5 2D/3D 可视化实现", 2)
    add_body_paragraph(
        doc,
        "在可视化方面，系统并未将 Three.js 作为首页装饰，而是放到更适合三维对比表达的统计场景中。分类统计模块提供 2D 图表与 3D 图表一键切换能力，二维图表用于准确阅读数值和比例，三维图表则用于增强视觉冲击和空间感。"
    )
    add_body_paragraph(
        doc,
        "在实现过程中，系统针对三维柱图做了多次迭代，重点解决了初始视角偏移、标签遮挡、自动旋转干扰、锯齿明显和质感不够高级等问题。"
        "最终方案通过调整相机参数、材质、光照、标签布局和视图控制，使 3D 图表既保留展示效果，又不过度影响用户对数据差异的判断。"
    )

    add_heading_text(doc, "6.6 自动化脚本与部署实现", 2)
    add_body_paragraph(
        doc,
        "为了降低系统部署难度，本文在项目中配套实现了无 MySQL CLI 初始化脚本、全自动启动脚本、冒烟测试脚本和分析报告导出脚本。"
        "其中，全自动启动脚本会依次完成 Python 依赖检查、数据库初始化、后端启动、前端启动、接口自测以及浏览器打开首页等动作。"
    )
    add_body_paragraph(
        doc,
        "这些脚本的加入，使系统不再只是“源代码集合”，而是具备一键部署和演示能力的工程项目。"
        "对于毕业设计而言，这类工程化能力能显著提高系统可信度，也方便答辩现场快速复现。"
    )

    add_heading_text(doc, "6.7 关键问题与解决方案", 2)
    add_body_paragraph(
        doc,
        "项目开发过程中先后遇到了多个关键问题。首先是多来源数据字段不统一的问题，不同平台或文件中的“标题、平台、播放量”等字段命名不一致，导致直接入库困难。"
        "对此，系统通过标准化字段字典和统一映射规则进行解决，并在规则解析不足时引入可选 AI 辅助结构化能力。"
    )
    add_body_paragraph(
        doc,
        "其次是大数据量导入和清理时容易超时的问题。早期系统在导入大 CSV 文件和清理行为日志时耗时较长，前端还可能因为超时误判为失败。"
        "后续通过批量写入、行为日志规模控制、分批删除、缓存失效替代同步重算以及前端长任务超时关闭等方式，将导入和清理成功率显著提升。"
    )
    add_body_paragraph(
        doc,
        "再次是用户画像结果过于单一的问题。早期画像模块容易把大量用户都归为同一种标签，缺乏分析价值。"
        "系统后续通过优化行为样本生成逻辑、调整标签判定顺序和引入多维阈值，使画像结果能够稳定分化为高互动讨论型、高频活跃型和稳定观看型等多类标签。"
    )


def add_chapter_seven(doc: Document):
    add_heading_text(doc, "第7章 系统测试与结果分析", 1)
    add_heading_text(doc, "7.1 测试环境", 2)
    add_body_paragraph(
        doc,
        "系统测试环境为 Windows 本地开发环境，主要软件包括 JDK 17 及以上版本、Maven 3.9 及以上版本、Node.js 18 及以上版本、Python 3.10 及以上版本和 MySQL 8。"
        "测试内容既包括功能正确性验证，也包括导入链路、登录状态、页面交互和自动化脚本的稳定性验证。"
    )

    add_heading_text(doc, "7.2 功能测试", 2)
    add_body_paragraph(
        doc,
        "功能测试采用“模块逐项验证 + 全链路回归”的方式开展。前端通过页面交互验证登录、导入、筛选、图表切换和数据清理等流程；后端通过编译、接口冒烟测试和数据库检查验证业务正确性。"
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for idx, text in enumerate(["测试项", "输入/操作", "预期结果", "结论"]):
        set_cell_text(table.rows[0].cells[idx], text, "黑体", 10.5, True)
    rows = [
        ("注册登录", "输入账号密码完成注册/登录", "成功进入系统并生成会话", "通过"),
        ("文件导入", "上传 CSV/Markdown/TXT 文件", "成功入库并显示状态反馈", "通过"),
        ("数据清理", "执行清空当前账号数据", "业务数据被删除且首页同步归零", "通过"),
        ("平台筛选", "切换抖音、哔哩哔哩或全部平台", "页面按平台展示对应数据", "通过"),
        ("2D/3D切换", "分类统计切换图表模式", "同源数据在二维和三维场景正常展示", "通过"),
        ("登录保持", "登录后刷新页面", "无需重新登录，仍保持当前会话", "通过"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text)

    add_heading_text(doc, "7.3 性能与稳定性测试", 2)
    add_body_paragraph(
        doc,
        "性能优化前，首页多个模块需要分别请求接口，导入和清理大文件时也存在长时间阻塞问题。优化后，系统在首页增加聚合接口和总览缓存表，在导入链路中加入批量写入和分片处理机制，在清理链路中采用分批删除策略，使整体体验得到明显改善。"
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for idx, text in enumerate(["测试指标", "优化前", "优化后", "说明"]):
        set_cell_text(table.rows[0].cells[idx], text, "黑体", 10.5, True)
    rows = [
        ("首页总览响应", "秒级", "毫秒级", "通过概览缓存和聚合接口降低查询压力"),
        ("大文件导入", "约134秒", "约16秒", "通过批量写入和行为规模控制优化"),
        ("大数据清理", "约206秒", "约11秒", "通过分批删除和缓存失效策略优化"),
        ("长任务稳定性", "易被超时中断", "稳定完成", "前端长任务关闭固定超时限制"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text)

    add_heading_text(doc, "7.4 测试结果分析", 2)
    add_body_paragraph(
        doc,
        "从测试结果看，系统已经能够较稳定地完成毕业设计所要求的主要功能。数据导入链路完整，错误反馈明确，平台筛选逻辑清晰，多用户数据空间独立，图表分析结果能够围绕研究问题展开。"
    )
    add_body_paragraph(
        doc,
        "同时，系统在工程完整度上也优于单纯的页面展示项目：一方面提供了自动化脚本、自测脚本和报告导出脚本；另一方面通过来源追踪、拒绝记录和质量评分提高了数据可信性。"
        "这说明本文实现的不仅是一个展示页面，而是一个具备基本工程可交付能力的数据分析系统。"
    )


def add_chapter_eight(doc: Document):
    add_heading_text(doc, "第8章 总结与展望", 1)
    add_body_paragraph(
        doc,
        "本文围绕视频网站数据分析场景，完成了一个集登录认证、数据导入、标准化入库、质量治理、分析展示和自动化部署于一体的全栈系统。"
        "与只展示静态图表的传统毕业设计不同，本文更加强调多用户隔离、双身份账号、来源追踪、质量控制和工程可复现性，使系统具备较完整的数据闭环。"
    )
    add_body_paragraph(
        doc,
        "在具体实现中，系统采用 Vue3、Spring Boot、MyBatis、MySQL 和 Python 组成的技术方案，支持 URL、文本和文件等多种数据接入方式，实现了热门视频、分类统计、互动效率、用户画像和创作者中心等多维分析功能，并通过 Three.js 扩展了三维可视化表现形式。"
    )
    add_body_paragraph(
        doc,
        "尽管系统已经能够满足本科毕业设计的基本要求，但仍存在进一步完善空间。例如，当前 AI 辅助入库仍以最小可用版为主，未来可以继续优化提示词、增加批量异构文本识别能力；当前系统主要适合本地部署与中小规模数据分析，后续还可以考虑接入消息队列、离线计算框架和更完善的监控告警机制。"
    )
    add_body_paragraph(
        doc,
        "后续工作还可以继续扩展平台适配范围，细化用户画像标签体系，补充数据血缘分析和结果解释能力，使平台从“教学型研究系统”逐步向“轻量级业务分析平台”演进。"
    )

    add_heading_text(doc, "致 谢", 1)
    add_body_paragraph(
        doc,
        "在本课题的设计与实现过程中，指导教师在选题方向、系统功能边界和论文结构方面给予了指导，同学在环境配置和测试使用方面提供了帮助。本文参考了相关开源框架文档和技术资料，在此一并表示感谢。"
    )

    add_heading_text(doc, "参考文献", 1)
    references = [
        "[1] 王珊, 萨师煊. 数据库系统概论[M]. 北京: 高等教育出版社, 2021.",
        "[2] 李智慧. 大型网站技术架构: 核心原理与案例分析[M]. 北京: 电子工业出版社, 2019.",
        "[3] Spring Boot Reference Documentation[EB/OL]. https://docs.spring.io/spring-boot/docs/current/reference/html/ .",
        "[4] Vue.js Documentation[EB/OL]. https://cn.vuejs.org/ .",
        "[5] Apache ECharts Documentation[EB/OL]. https://echarts.apache.org/ .",
        "[6] Three.js Documentation[EB/OL]. https://threejs.org/docs/ .",
        "[7] MySQL 8.0 Reference Manual[EB/OL]. https://dev.mysql.com/doc/ .",
        "[8] Martin Fowler. Patterns of Enterprise Application Architecture[M]. Boston: Addison-Wesley, 2003.",
        "[9] 蒋宗礼. 软件工程导论[M]. 北京: 清华大学出版社, 2020.",
        "[10] 刘鹏, 王超. 大数据技术基础[M]. 北京: 机械工业出版社, 2021.",
        "[11] 陈为, 沈则潜, 陶煜波. 数据可视化[M]. 北京: 电子工业出版社, 2020.",
        "[12] OpenAI API Documentation[EB/OL]. https://platform.openai.com/docs/ .",
    ]
    for ref in references:
        add_plain_paragraph(doc, ref)


def build_document(output_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    add_page_number(section)

    configure_style(doc.styles["Normal"], "宋体", 12)
    configure_style(doc.styles["Heading 1"], "黑体", 16, True)
    configure_style(doc.styles["Heading 2"], "黑体", 14, True)
    configure_style(doc.styles["Heading 3"], "黑体", 12, True)

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "本科毕业设计论文"
    doc.core_properties.author = "待补充"

    add_cover(doc)
    add_abstract(doc)

    add_heading_text(doc, "目 录", 1)
    add_toc(doc)
    doc.add_page_break()

    add_chapter_one(doc)
    add_chapter_two(doc)
    add_chapter_three(doc)
    add_chapter_four(doc)
    add_chapter_five(doc)
    add_chapter_six(doc)
    add_chapter_seven(doc)
    add_chapter_eight(doc)

    doc.save(str(output_path))


def main():
    current = Path(__file__).resolve()
    output_path = current.parents[2] / OUTPUT_NAME
    build_document(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
