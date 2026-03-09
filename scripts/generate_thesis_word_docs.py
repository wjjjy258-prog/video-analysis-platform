from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "瀹嬩綋"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "瀹嬩綋")
    normal.font.size = Pt(11)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "榛戜綋"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "榛戜綋")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_meta(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def generate_doc1(output_dir: Path) -> Path:
    doc = Document()
    setup_styles(doc)

    add_title(doc, "瑙嗛缃戠珯鏁版嵁鍒嗘瀽骞冲彴寤鸿鍏ㄨ繃绋嬩笌闂瑙ｅ喅鎶ュ憡")
    add_meta(
        doc,
        [
            "椤圭洰鍚嶇О锛氳棰戠綉绔欐暟鎹垎鏋愬钩鍙帮紙Video Analysis Platform锛?,
            f"鏂囨。鏃ユ湡锛歿datetime.now().strftime('%Y-%m-%d')}",
            "鏂囨。鐢ㄩ€旓細姣曚笟璁捐杩囩▼鎬ф潗鏂欙紙浣撶幇閫愭寤鸿涓庡伐绋嬫€濊矾锛?,
        ],
    )

    doc.add_heading("1. 椤圭洰鑳屾櫙涓庡缓璁剧洰鏍?, level=1)
    doc.add_paragraph(
        "鏈」鐩潰鍚戞姈闊炽€丅ilibili 绛夎棰戞暟鎹垎鏋愬満鏅紝鐩爣涓嶆槸鍙仛涓€涓€滃彲灞曠ず椤甸潰鈥濓紝鑰屾槸鏋勫缓涓€鏉″畬鏁寸殑鏁版嵁闂幆锛?
        "鏁版嵁閲囬泦锛圲RL/鏂囨湰/鏂囦欢锛? 鏁版嵁鍏ュ簱锛堝幓閲嶃€佽拷韪€佽川閲忚瘎鍒嗭級- 鏁版嵁鍒嗘瀽锛堜簰鍔ㄦ晥鐜囥€佺敤鎴风敾鍍忋€佸钩鍙板姣旓級- 鍙鍖栧憟鐜般€?
        "寤鸿杩囩▼涓潥鎸佲€滃厛鍙敤锛屽啀瀹屽杽锛涘厛闂幆锛屽啀浼樺寲锛涘厛绋冲畾锛屽啀鎵╁睍鈥濈殑宸ョ▼绛栫暐銆?
    )
    add_bullets(
        doc,
        [
            "鐩爣1锛氬畬鎴愬墠绔€佸悗绔€佹暟鎹簱涓€浣撳寲绯荤粺锛屾敮鎸佹湰鍦扮洿鎺ラ儴缃层€?,
            "鐩爣2锛氭敮鎸佸鏉ユ簮鏁版嵁鎺ュ叆锛屼繚璇佹暟鎹悎娉曟€с€佸彲杩借釜鎬у拰鍙В閲婃€с€?,
            "鐩爣3锛氬垎鏋愮粨鏋滆鏈嶅姟浜庣爺绌堕棶棰橈紝閬垮厤浠呭仛鈥滃浘琛ㄥ爢鍙犫€濄€?,
            "鐩爣4锛氬舰鎴愬彲澶嶇幇鐨勪竴閿剼鏈笌鏂囨。锛岄檷浣庨儴缃蹭笌婕旂ず鎴愭湰銆?,
        ],
    )

    doc.add_heading("2. 鎬讳綋寤鸿鏂规硶璁?, level=1)
    doc.add_paragraph(
        "椤圭洰閲囩敤鈥滃垎闃舵澧為噺浜や粯鈥濇柟寮忔帹杩涳紝姣忎釜闃舵閮藉舰鎴愬彲楠岃瘉缁撴灉锛岄伩鍏嶅ぇ鑰屽叏涓€娆℃€у紑鍙戝鑷撮闄╁け鎺с€?
    )
    add_numbered(
        doc,
        [
            "闃舵A锛氱‘瀹氭妧鏈爤涓庢渶灏忕郴缁熼鏋讹紙鍓嶇璺敱銆佸悗绔仴搴锋帴鍙ｃ€佹暟鎹簱鍒濆鍖栵級銆?,
            "闃舵B锛氭墦閫氭牳蹇冨垎鏋愰〉闈紙鐑棬瑙嗛銆佸垎绫荤粺璁°€佷簰鍔ㄦ晥鐜囥€佺敤鎴风敾鍍忥級銆?,
            "闃舵C锛氳ˉ榻愭暟鎹噰闆嗕笌鍏ュ簱鑳藉姏锛圲RL銆佹枃鏈€佹枃浠跺鍏ワ級銆?,
            "闃舵D锛氬畬鍠勬暟鎹川閲忎笌杩借釜锛坉edupe_key銆乮mport_job銆乻ource_trace锛夈€?,
            "闃舵E锛氭彁鍗囧彲瑙嗗寲涓庝氦浜掞紙2D/3D鍒囨崲銆乀hree.js缁嗚妭浼樺寲锛夈€?,
            "闃舵F锛氳剼鏈寲閮ㄧ讲涓庤嚜娴嬶紙鏃犲懡浠よ鍒濆鍖栥€佷竴閿叏鑷姩鍚姩銆佹姤鍛婂鍑猴級銆?,
            "闃舵G锛氶鏍肩粺涓€涓庡姛鑳借鍓紙鍒犻櫎涓嶉渶瑕佸姛鑳斤紝淇濇寔绯荤粺涓€鑷存€э級銆?,
        ],
    )

    doc.add_heading("3. 鍒嗛樁娈靛缓璁捐繃绋嬶紙涓€姝ヤ竴姝ュ畬鎴愶級", level=1)

    doc.add_heading("3.1 闃舵A锛氭惌寤哄熀纭€楠ㄦ灦", level=2)
    doc.add_paragraph(
        "鍏堣惤鍦板彲杩愯鐨勪笁灞傛灦鏋勶細MySQL 瀛樺偍銆丼pring Boot + MyBatis 鎻愪緵 REST銆乂ue3 鎻愪緵鍓嶇璺敱涓庨〉闈㈠鍣ㄣ€?
    )
    add_bullets(
        doc,
        [
            "瀹屾垚鏁版嵁搴撳垵濮嬪寲鑴氭湰 `database/init.sql`锛屽畾涔?video/user/comment/user_behavior/video_statistics 绛夊熀纭€琛ㄣ€?,
            "瀹屾垚鍚庣鍩虹鎺ュ彛锛?health銆?video/overview銆?video/hot 绛夛級骞舵墦閫?MyBatis 鏄犲皠銆?,
            "瀹屾垚鍓嶇瀵艰埅涓庡熀纭€椤甸潰缁撴瀯锛屽厛淇濊瘉鎺ュ彛鍙覆鏌擄紝鍐嶅仛瑙嗚浼樺寲銆?,
        ],
    )
    doc.add_paragraph("闃舵缁撹锛氱郴缁熷叿澶団€滃彲鍚姩銆佸彲璁块棶銆佸彲鐪嬪埌鏁版嵁鈥濈殑鏈€浣庡彲鐢ㄨ兘鍔涖€?)

    doc.add_heading("3.2 闃舵B锛氬垎鏋愬姛鑳芥垚鍨?, level=2)
    doc.add_paragraph(
        "鍥寸粫鈥滆棰戞暟鎹垎鏋愨€濅富绾匡紝閫愭瀹炵幇鏍稿績鍒嗘瀽椤甸潰锛屽舰鎴愮爺绌跺彲鐢ㄧ殑鎸囨爣浣撶郴銆?
    )
    add_bullets(
        doc,
        [
            "鐑棬瑙嗛锛氭寜鎾斁/鐐硅禐鎺掑簭锛屾敮鎸佸钩鍙板瓧娈靛睍绀恒€?,
            "鍒嗙被缁熻锛氭寜鍒嗙被鑱氬悎鎾斁銆佺偣璧炪€佽瘎璁猴紝鏀寔 2D/3D 鍥惧垏鎹€?,
            "浜掑姩鏁堢巼锛氫粠鍗曠函鈥滄挱鏀捐秼鍔库€濊皟鏁翠负鏇存湁浠峰€肩殑浜掑姩鐜囧垎鏋愩€?,
            "鐢ㄦ埛鐢诲儚锛氳绠楁椿璺冨害銆佸叴瓒ｅ箍搴︺€佸钩鍙板亸濂斤紝骞跺紩鍏ョ敾鍍忔爣绛俱€?,
        ],
    )
    doc.add_paragraph(
        "闃舵缁撹锛氬畬鎴愨€滀粠鏁版嵁鍒板垎鏋愯鍥锯€濈殑鏍稿績鍔熻兘锛岄〉闈㈠凡鍏峰鐮旂┒灞曠ず浠峰€笺€?
    )

    doc.add_heading("3.3 闃舵C锛氭暟鎹噰闆嗘帴鍏?, level=2)
    doc.add_paragraph(
        "瑙ｅ喅鈥滄暟鎹簱鏈夌粨鏋勪絾鏃犵湡瀹炴暟鎹€濈殑闂锛屾瀯寤洪噰闆嗕笌瀵煎叆鍏ュ彛锛屽苟鍏奸【鍚堣杈圭晫銆?
    )
    add_bullets(
        doc,
        [
            "鏂板骞冲彴鍐呪€滄暟鎹噰闆嗏€濋〉闈紝鏀寔 URL 閲囬泦銆佹ā鎷熼噰闆嗐€佹枃鏈鍏ャ€佹枃浠跺鍏ャ€?,
            "鏂囦欢瀵煎叆鏀寔 `.md/.txt/.markdown/.log/.csv`锛屽苟鍋氭牸寮忚В鏋愪笌瀛楁鎶藉彇銆?,
            "鏂板 URL 妯℃澘鐢熸垚涓庢牸寮忔牎楠岋紝闄嶄綆鎵归噺瀵煎叆鍑洪敊鐜囥€?,
            "寮鸿皟楂橀闄╂搷浣滅殑鎺堟潈纭锛岄伩鍏嶉粯璁ら潪娉曟姄鍙栨祦绋嬨€?,
        ],
    )
    doc.add_paragraph("闃舵缁撹锛氱郴缁熷叿澶囨寔缁繘鏁拌兘鍔涳紝涓嶅啀渚濊禆鎵嬪伐鎻掑叆鏍蜂緥鏁版嵁銆?)

    doc.add_heading("3.4 闃舵D锛氭暟鎹彲淇￠棴鐜紙鍏抽敭鍗囩骇锛?, level=2)
    doc.add_paragraph(
        "鍦ㄧ湡瀹炲伐绋嬩腑锛屾暟鎹彲杩借釜涓庡幓閲嶆瘮鈥滃鍑犱釜鍥锯€濇洿鍏抽敭銆傛闃舵閲嶇偣瑙ｅ喅鈥滈噸澶嶅鍏ャ€佹潵婧愪笉鏄庛€佽川閲忎笉鍙帶鈥濈殑鏍稿績椋庨櫓銆?
    )
    add_bullets(
        doc,
        [
            "video 琛ㄦ柊澧?`dedupe_key/source_url/import_type/source_file/import_time/data_quality_score/import_job_id`銆?,
            "鏂板 `import_job` 琛ㄨ褰曟瘡娆″鍏ヤ换鍔★紝鏀寔瀹¤涓庡洖婧€?,
            "瀵煎叆閫昏緫鎸?dedupe_key 鍘婚噸骞跺悎骞讹紝閬垮厤閲嶅瑙嗛瀵艰嚧缁熻澶辩湡銆?,
            "鏂板璐ㄩ噺璇勫垎瑙勫垯锛岄噺鍖栨暟鎹畬鏁存€э紙鏍囬銆佷綔鑰呫€佸钩鍙般€佷簰鍔ㄩ噺銆佹椂闂寸瓑锛夈€?,
            "涓烘棫搴撳鍔犺嚜鍔ㄥ崌绾ц兘鍔涳紙SchemaUpgradeRunner + upgrade SQL锛夈€?,
        ],
    )
    doc.add_paragraph("闃舵缁撹锛氱郴缁熶粠鈥滆兘灞曠ず鈥濆崌绾у埌鈥滅粨鏋滃彲淇°€佹潵婧愬彲杩借釜鈥濄€?)

    doc.add_heading("3.5 闃舵E锛氬彲瑙嗗寲涓庝綋楠屼紭鍖?, level=2)
    doc.add_paragraph(
        "鍥寸粫鐢ㄦ埛鍙嶉杩涜鍓嶇杩唬锛岀洰鏍囨槸鈥滅畝绾︺€佸ぇ姘斻€佹竻鏅般€佸彲璇烩€濄€?
    )
    add_bullets(
        doc,
        [
            "鏁翠綋椋庢牸鏀逛负绠€娲佷俊鎭寲璁捐锛屼紭鍖栧鑸€佸眰绾с€佺暀鐧戒笌鏁版嵁鍗″竷灞€銆?,
            "Three.js 浠庨椤佃楗版€у睍绀鸿縼绉诲埌鏇存湁鎰忎箟鐨勫姣斿浘鍦烘櫙銆?,
            "瀹炵幇 2D/3D 涓€閿垏鎹紝璁╃敤鎴峰彲鎸夊満鏅€夋嫨琛ㄨ揪鏂瑰紡銆?,
            "淇閿娇銆佽瑙掑亸绉汇€佹爣绛鹃伄鎸°€佽嚜鍔ㄦ棆杞共鎵扮瓑鍙鍖栭棶棰樸€?,
            "绉婚櫎涓嶉渶瑕佸姛鑳斤紙濡傛渶缁堝垹闄も€滅瓟杈╂ā寮忊€濓級锛屼繚鎸佺郴缁熶竴鑷存€с€?
        ],
    )

    doc.add_heading("3.6 闃舵F锛氶儴缃茶嚜鍔ㄥ寲涓庡伐绋嬫敹鍙?, level=2)
    add_bullets(
        doc,
        [
            "鎻愪緵鈥滄棤鍛戒护琛岀増鈥濊剼鏈紙鍒濆鍖?+ 鑷祴锛夈€?,
            "鎻愪緵鈥滃叏鑷姩鐗堚€濊剼鏈紙鍚姩鍚庣銆佸墠绔€佹墦寮€骞冲彴椤甸潰锛夈€?,
            "淇 PowerShell 缂栫爜瀵艰嚧鐨勫瓧绗︿覆缁堟绗﹂敊璇€?,
            "鑷祴鑴氭湰澧炲姞瓒呮椂閲嶈瘯涓庡彲閫夋帴鍙ｈ鍛婃満鍒讹紝閬垮厤璇垽銆?,
            "鏂板 `export_analysis_report.py` 瀵煎嚭 CSV/Markdown 鎶ュ憡鍒?`analysis/reports`銆?,
        ],
    )
    doc.add_paragraph("闃舵缁撹锛氶」鐩叿澶囦氦浠樺睘鎬э紝鍙鐜般€佸彲婕旂ず銆佸彲鍥炲綊楠岃瘉銆?)

    doc.add_heading("4. 鍏抽敭闅鹃涓庤В鍐崇瓥鐣?, level=1)
    add_table(
        doc,
        ["闅鹃", "琛ㄧ幇", "鏍瑰洜", "瑙ｅ喅绛栫暐", "鏁堟灉"],
        [
            [
                "MySQL 瀹㈡埛绔棤娉曚娇鐢?,
                "鍛戒护琛岀櫥褰曟姤 ODBC 鐢ㄦ埛閿欒",
                "鐜鍙橀噺/榛樿瀹㈡埛绔厤缃紓甯?,
                "鏀圭敤 Python 鐩磋繛鍒濆鍖栵紙mysql-connector锛? Navicat 鍙鍖栨搷浣?,
                "涓嶄緷璧?mysql CLI 涔熷彲瀹屾垚鍒濆鍖栦笌楠岃瘉",
            ],
            [
                "鑴氭湰鍚姩鎶?PowerShell 瑙ｆ瀽閿欒",
                "瀛楃涓茬己灏戠粓姝㈢銆佺己鍙虫嫭鍙?,
                "鏂囦欢缂栫爜鍜屼腑鏂囧瓧绗︿覆鍦ㄨ剼鏈腑鎹熷潖",
                "鑴氭湰缁熶竴鏀逛负 UTF-8锛岃嫳鏂囬敊璇彁绀猴紝閲嶆柊鏁寸悊 start_* 鑴氭湰",
                "鍏ㄨ嚜鍔ㄨ剼鏈ǔ瀹氳繍琛?,
            ],
            [
                "鍚庣鍚姩澶辫触锛圡yBatis XML锛?,
                "sqlSessionFactory 鍒濆鍖栧紓甯?,
                "XML 鐗规畩瀛楃涓庤娉曢敊璇?,
                "閫愯瀹氫綅 mapper XML锛屼慨澶嶄笉鍚堟硶瀛楃涓?SQL 鐗囨",
                "鍚庣缂栬瘧骞跺惎鍔ㄦ仮澶嶆甯?,
            ],
            [
                "鑷祴鎺ュ彛瓒呮椂",
                "Smoke test ReadTimeout",
                "鍚姩鍚庨娆℃煡璇㈣€楁椂澶с€佽秴鏃惰缃繃浣?,
                "鎻愰珮瓒呮椂銆佸鍔犻噸璇曘€佸尯鍒嗕弗鏍?闈炰弗鏍兼ā寮?,
                "鑷姩鍖栨祦绋嬫姉鎶栧姩鑳藉姏鎻愬崌",
            ],
            [
                "椤甸潰涓枃涔辩爜",
                "鏍囬涓庢枃妗堟樉绀哄紓甯?,
                "鏂囦欢缂栫爜涓庣粓绔紪鐮佹贩鏉?,
                "鍓嶇婧愮爜缁熶竴 UTF-8锛岄€愰〉淇鏂囨",
                "椤甸潰鏄剧ず鎭㈠姝ｅ父",
            ],
            [
                "Three.js 鍥捐〃涓嶅疄鐢?,
                "鏃嬭浆褰卞搷闃呰銆佺己灏戞爣绛俱€佽川鎰熶笉瓒?,
                "鍒濈増鍋忔紨绀洪鏍硷紝寮变笟鍔¤〃杈?,
                "鍔犲叆鏍囩銆侀噸缃瑙掋€佸彲寮€鍏虫棆杞€佹潗璐?鍏夌収/鎶楅敮榻夸紭鍖?,
                "3D鍥惧叿澶囧垎鏋愯緟鍔╀环鍊?,
            ],
            [
                "鏁版嵁閲嶅涓庢潵婧愪笉娓?,
                "瀵煎叆鍚庢暟閲忓紓甯搞€侀毦浠ヨ拷婧?,
                "缂轰箯缁熶竴鍘婚噸閿拰瀵煎叆浠诲姟鏃ュ織",
                "寮曞叆 dedupe_key + import_job + source_trace + quality_score",
                "缁撴灉鍙俊鎬ф樉钁楁彁鍗?,
            ],
        ],
    )

    doc.add_heading("5. 璐ㄩ噺淇濋殰涓庨獙璇佹満鍒?, level=1)
    add_bullets(
        doc,
        [
            "缂栬瘧楠岃瘉锛氬悗绔?`mvn -DskipTests compile`锛屽墠绔?`npm run build`銆?,
            "杩愯楠岃瘉锛歚/health`銆乣/video/*` 鍏抽敭鎺ュ彛鍐掔儫妫€鏌ャ€?,
            "鏁版嵁楠岃瘉锛氭鏌?video/user/user_behavior 鏁伴噺銆佸鍏ヤ换鍔¤褰曘€佹潵婧愯拷韪褰曘€?,
            "鑴氭湰楠岃瘉锛氭棤鍛戒护琛屽垵濮嬪寲鑴氭湰銆佸叏鑷姩鍚姩鑴氭湰銆佹姤鍛婂鍑鸿剼鏈€?,
            "鏂囨。楠岃瘉锛氶儴缃叉墜鍐屻€佺埇铏鏄庝笌鍔熻兘涓€鑷存€у悓姝ャ€?,
        ],
    )

    doc.add_heading("6. 杩囩▼澶嶇洏涓庣粡楠屾€荤粨", level=1)
    doc.add_paragraph(
        "鏈」鐩渶缁堣兘澶熷畬鎴愶紝鏄洜涓哄缁堝潥鎸佲€滀互闂幆涓轰腑蹇冣€濈殑杩唬鏂瑰紡銆傛瘡娆℃敼鍔ㄩ兘鍥寸粫涓€涓槑纭棶棰樻帹杩涳紝"
        "骞剁敤浠ｇ爜銆佽剼鏈拰鏂囨。涓夋潯绾垮悓姝ユ敹鍙ｃ€傚疄璺佃瘉鏄庯紝姣曚笟璁捐绫荤郴缁熷鏋滃彧杩芥眰椤甸潰鏁堟灉锛屽緢闅剧粡鍙楃瓟杈╄拷闂紱"
        "鑰屽綋绯荤粺鍏峰鏉ユ簮杩借釜銆佹暟鎹川閲忔帶鍒躲€佽嚜鍔ㄥ寲閮ㄧ讲涓庡彲澶嶇幇瀹為獙璺緞鏃讹紝鏁翠綋鍙俊搴﹀拰宸ョ▼瀹屾垚搴︿細鏄庢樉鎻愰珮銆?
    )
    add_bullets(
        doc,
        [
            "缁忛獙1锛氬厛瑙ｅ喅鈥滆兘璺戦€氣€濓紝鍐嶈В鍐斥€滆窇寰楀ソ鈥濄€?,
            "缁忛獙2锛氳剼鏈笌鏂囨。涓嶆槸闄勫睘鍝侊紝鑰屾槸浜や粯鐨勪竴閮ㄥ垎銆?,
            "缁忛獙3锛氭暟鎹垎鏋愮郴缁熸渶鎬曗€滃亣鏁版嵁缁撹鈥濓紝蹇呴』鍋氭潵婧愪笌璐ㄩ噺娌荤悊銆?,
            "缁忛獙4锛氬姛鑳戒笉鍦ㄥ锛屽湪浜庢槸鍚︽敮鎾戠爺绌堕棶棰樹笌涓氬姟瑙ｉ噴銆?,
        ],
    )

    path = output_dir / "骞冲彴寤鸿鍏ㄨ繃绋嬩笌闂瑙ｅ喅鎶ュ憡.docx"
    doc.save(path)
    return path


def generate_doc2(output_dir: Path) -> Path:
    doc = Document()
    setup_styles(doc)

    add_title(doc, "瑙嗛缃戠珯鏁版嵁鍒嗘瀽骞冲彴鏁版嵁搴撹璁¤鏄庝功")
    add_meta(
        doc,
        [
            "鏁版嵁搴擄細MySQL 8.x",
            f"鏂囨。鏃ユ湡锛歿datetime.now().strftime('%Y-%m-%d')}",
            "璁捐鐩爣锛氭敮鎾戦噰闆嗗叆搴撱€佸垎鏋愯绠椼€佺粨鏋滆拷韪拰鍙墿灞曟紨杩?,
        ],
    )

    doc.add_heading("1. 鏁版嵁搴撹璁＄洰鏍囦笌鍘熷垯", level=1)
    add_bullets(
        doc,
        [
            "瀹屾暣鎬э細瑕嗙洊瑙嗛銆佺敤鎴枫€佽涓恒€佺粺璁°€佸鍏ヤ换鍔″叏閾捐矾鏁版嵁銆?,
            "涓€鑷存€э細閫氳繃涓婚敭銆佸敮涓€閿€佺储寮曞拰瑙勫垯绾︽潫淇濊瘉鏁版嵁涓€鑷淬€?,
            "鍙拷韪細璁板綍鏉ユ簮 URL銆佸鍏ョ被鍨嬨€佹簮鏂囦欢銆佸鍏ユ椂闂淬€佷换鍔D銆?,
            "鍙墿灞曪細鏀寔鍚庣画澧炲姞骞冲彴銆佹爣绛俱€佹帹鑽愪笌瀹為獙妯″潡銆?,
            "鎬ц兘鍙帶锛氬楂橀鏌ヨ瀛楁寤虹珛绱㈠紩锛屽吋椤惧啓鍏ヤ笌鏌ヨ銆?,
        ],
    )

    doc.add_heading("2. 姒傚康缁撴瀯璁捐锛堝疄浣撲笌鍏崇郴锛?, level=1)
    doc.add_paragraph("鏍稿績瀹炰綋鍖呮嫭锛歏ideo銆乁ser銆丆omment銆乁serBehavior銆乂ideoStatistics銆乁serInterestResult銆両mportJob銆?)
    add_bullets(
        doc,
        [
            "Video 涓?UserBehavior锛氫竴瀵瑰锛堜竴涓棰戝搴斿鏉¤涓烘棩蹇楋級銆?,
            "User 涓?UserBehavior锛氫竴瀵瑰锛堜竴涓敤鎴峰搴斿鏉¤涓烘棩蹇楋級銆?,
            "Video 涓?Comment锛氫竴瀵瑰锛堜竴涓棰戝彲鏈夊鏉¤瘎璁猴級銆?,
            "Video 涓?VideoStatistics锛氫竴瀵瑰锛堟寜鏃ユ湡鑱氬悎缁熻锛夈€?,
            "User 涓?UserInterestResult锛氫竴瀵逛竴锛堟瘡涓敤鎴蜂竴鏉＄敾鍍忚仛绫荤粨鏋滐級銆?,
            "ImportJob 涓?Video锛氫竴瀵瑰锛堜竴娆″鍏ヤ换鍔″搴斿鏉¤棰戝叆搴撹褰曪級銆?,
        ],
    )

    doc.add_heading("3. 閫昏緫缁撴瀯璁捐锛堣〃缁撴瀯璇﹁В锛?, level=1)

    doc.add_heading("3.1 video锛堣棰戜富琛級", level=2)
    add_table(
        doc,
        ["瀛楁", "绫诲瀷", "绾︽潫", "璇存槑"],
        [
            ["id", "BIGINT", "PK, AUTO_INCREMENT", "瑙嗛涓婚敭ID"],
            ["dedupe_key", "VARCHAR(128)", "UNIQUE", "鍘婚噸閿紝閬垮厤閲嶅瀵煎叆"],
            ["title", "VARCHAR(255)", "NOT NULL", "瑙嗛鏍囬"],
            ["author", "VARCHAR(100)", "NOT NULL", "浣滆€呭悕绉?],
            ["source_platform", "VARCHAR(32)", "NOT NULL", "鏉ユ簮骞冲彴锛坆ilibili/douyin/seed绛夛級"],
            ["source_url", "VARCHAR(1000)", "NULL", "鏉ユ簮閾炬帴锛堟爣鍑嗗寲锛?],
            ["category", "VARCHAR(60)", "NOT NULL", "瑙嗛鍒嗙被"],
            ["play_count", "BIGINT", "NOT NULL", "鎾斁閲?],
            ["like_count", "BIGINT", "NOT NULL", "鐐硅禐閲?],
            ["comment_count", "BIGINT", "NOT NULL", "璇勮閲?],
            ["publish_time", "DATETIME", "NOT NULL", "鍙戝竷鏃堕棿"],
            ["import_type", "VARCHAR(32)", "NOT NULL", "瀵煎叆鏂瑰紡锛坈rawler_url/file_import绛夛級"],
            ["source_file", "VARCHAR(260)", "NULL", "婧愭枃浠跺悕"],
            ["import_time", "DATETIME", "NOT NULL", "瀵煎叆鏃堕棿"],
            ["data_quality_score", "DECIMAL(5,2)", "NOT NULL", "鏁版嵁璐ㄩ噺璇勫垎(0-100)"],
            ["import_job_id", "BIGINT", "NULL", "瀵煎叆浠诲姟ID"],
            ["created_at", "TIMESTAMP", "NOT NULL", "鍒涘缓鏃堕棿"],
        ],
    )
    doc.add_paragraph("绱㈠紩锛歚uk_video_dedupe(dedupe_key)`銆乣idx_video_import_time(import_time)`銆?)

    doc.add_heading("3.2 user锛堢敤鎴疯〃锛?, level=2)
    add_table(
        doc,
        ["瀛楁", "绫诲瀷", "绾︽潫", "璇存槑"],
        [
            ["user_id", "BIGINT", "PK, AUTO_INCREMENT", "鐢ㄦ埛涓婚敭"],
            ["user_name", "VARCHAR(100)", "NOT NULL", "鐢ㄦ埛鍚?],
            ["fans", "BIGINT", "NOT NULL", "绮変笣鏁?],
            ["follow", "BIGINT", "NOT NULL", "鍏虫敞鏁?],
            ["level", "INT", "NOT NULL", "绛夌骇"],
            ["created_at", "TIMESTAMP", "NOT NULL", "鍒涘缓鏃堕棿"],
        ],
    )

    doc.add_heading("3.3 comment锛堣瘎璁鸿〃锛?, level=2)
    add_table(
        doc,
        ["瀛楁", "绫诲瀷", "绾︽潫", "璇存槑"],
        [
            ["comment_id", "BIGINT", "PK, AUTO_INCREMENT", "璇勮涓婚敭"],
            ["video_id", "BIGINT", "NOT NULL, INDEX", "鎵€灞炶棰慖D"],
            ["user_id", "BIGINT", "NOT NULL, INDEX", "璇勮鐢ㄦ埛ID"],
            ["content", "VARCHAR(1000)", "NOT NULL", "璇勮鍐呭"],
            ["like_count", "BIGINT", "NOT NULL", "璇勮鐐硅禐鏁?],
            ["time", "DATETIME", "NOT NULL", "璇勮鏃堕棿"],
        ],
    )

    doc.add_heading("3.4 user_behavior锛堢敤鎴疯涓烘棩蹇楋級", level=2)
    add_table(
        doc,
        ["瀛楁", "绫诲瀷", "绾︽潫", "璇存槑"],
        [
            ["id", "BIGINT", "PK, AUTO_INCREMENT", "琛屼负鏃ュ織涓婚敭"],
            ["user_id", "BIGINT", "NOT NULL, INDEX", "鐢ㄦ埛ID"],
            ["video_id", "BIGINT", "NOT NULL, INDEX", "瑙嗛ID"],
            ["action", "VARCHAR(32)", "NOT NULL", "琛屼负绫诲瀷锛坧lay/like/comment锛?],
            ["time", "DATETIME", "NOT NULL, INDEX", "琛屼负鏃堕棿"],
        ],
    )

    doc.add_heading("3.5 video_statistics锛堣棰戞棩缁熻锛?, level=2)
    add_table(
        doc,
        ["瀛楁", "绫诲瀷", "绾︽潫", "璇存槑"],
        [
            ["id", "BIGINT", "PK, AUTO_INCREMENT", "涓婚敭"],
            ["video_id", "BIGINT", "NOT NULL", "瑙嗛ID"],
            ["stat_date", "DATE", "NOT NULL", "缁熻鏃ユ湡"],
            ["daily_play", "BIGINT", "NOT NULL", "鏃ユ挱鏀鹃噺"],
            ["daily_like", "BIGINT", "NOT NULL", "鏃ョ偣璧為噺"],
            ["daily_comment", "BIGINT", "NOT NULL", "鏃ヨ瘎璁洪噺"],
        ],
    )
    doc.add_paragraph("绾︽潫锛歚UNIQUE(video_id, stat_date)`锛岄槻姝㈠悓瑙嗛鍚屾棩鏈熼噸澶嶇粺璁°€?)

    doc.add_heading("3.6 user_interest_result锛堢敤鎴峰叴瓒ｈ仛绫荤粨鏋滐級", level=2)
    add_table(
        doc,
        ["瀛楁", "绫诲瀷", "绾︽潫", "璇存槑"],
        [
            ["id", "BIGINT", "PK, AUTO_INCREMENT", "涓婚敭"],
            ["user_id", "BIGINT", "UNIQUE, NOT NULL", "鐢ㄦ埛ID"],
            ["cluster_id", "INT", "NOT NULL", "鑱氱被缂栧彿"],
            ["cluster_label", "VARCHAR(64)", "NOT NULL", "鑱氱被鏍囩"],
            ["favorite_category", "VARCHAR(64)", "NOT NULL", "鍋忓ソ鍒嗙被"],
            ["updated_at", "DATETIME", "NOT NULL", "鏇存柊鏃堕棿"],
        ],
    )

    doc.add_heading("3.7 import_job锛堝鍏ヤ换鍔¤〃锛?, level=2)
    add_table(
        doc,
        ["瀛楁", "绫诲瀷", "绾︽潫", "璇存槑"],
        [
            ["id", "BIGINT", "PK, AUTO_INCREMENT", "浠诲姟ID"],
            ["import_type", "VARCHAR(32)", "NOT NULL", "瀵煎叆鏂瑰紡"],
            ["source_platform", "VARCHAR(32)", "NULL", "鏉ユ簮骞冲彴"],
            ["source_file", "VARCHAR(260)", "NULL", "鏉ユ簮鏂囦欢"],
            ["source_count", "INT", "NOT NULL", "婧愯褰曟暟"],
            ["success_count", "INT", "NOT NULL", "鎴愬姛鍏ュ簱鏁?],
            ["started_at", "DATETIME", "NOT NULL", "寮€濮嬫椂闂?],
            ["finished_at", "DATETIME", "NULL", "缁撴潫鏃堕棿"],
            ["status", "VARCHAR(16)", "NOT NULL", "浠诲姟鐘舵€侊紙RUNNING/SUCCESS/FAILED锛?],
            ["notes", "VARCHAR(500)", "NULL", "澶囨敞"],
            ["created_at", "TIMESTAMP", "NOT NULL", "鍒涘缓鏃堕棿"],
        ],
    )

    doc.add_heading("4. 鍏抽敭璁捐绛栫暐", level=1)
    doc.add_heading("4.1 鍘婚噸绛栫暐", level=2)
    doc.add_paragraph(
        "閲囩敤 dedupe_key 浣滀负璺ㄦ潵婧愬幓閲嶆牳蹇冮敭锛屾瀯鎴愰€昏緫涓猴細骞冲彴 + 鏍囧噯鍖朥RL 鎴?骞冲彴+鏍囬+浣滆€呫€?
        "鍏ュ簱鏃堕€氳繃 UPSERT 鍚堝苟浜掑姩鎸囨爣锛岄槻姝㈤噸澶嶅鍏ラ€犳垚缁熻鑶ㄨ儉銆?
    )
    doc.add_heading("4.2 鏁版嵁璐ㄩ噺璇勫垎绛栫暐", level=2)
    doc.add_paragraph(
        "鏍规嵁鏍囬銆佷綔鑰呫€佸钩鍙般€佸垎绫汇€佹挱鏀鹃噺銆佺偣璧為噺銆佽瘎璁洪噺銆佸彂甯冩椂闂淬€佹潵婧怳RL绛夊瓧娈靛畬鏁村害璁＄畻璐ㄩ噺鍒嗭紝"
        "骞跺鏄庢樉寮傚父鏁版嵁杩涜鎵ｅ垎锛岀敤浜庡悗缁竻娲椾笌鍙俊搴﹀垽鏂€?
    )
    doc.add_heading("4.3 瀵煎叆杩借釜绛栫暐", level=2)
    doc.add_paragraph(
        "姣忔瀵煎叆鍒涘缓 import_job 璁板綍锛岃棰戣褰曞啓鍏?import_job_id锛屽疄鐜扳€滄煇鏉℃暟鎹粠鍝噷鏉ャ€佷綍鏃跺叆搴撱€佺敱璋佸鍏モ€濆彲杩芥函銆?
    )

    doc.add_heading("5. 绱㈠紩涓庢€ц兘璁捐", level=1)
    add_bullets(
        doc,
        [
            "琛屼负鏃ュ織鎸?user_id/video_id/time 寤虹储寮曪紝鏀拺鐢ㄦ埛鐢诲儚涓庢椂搴忓垎鏋愩€?,
            "瑙嗛琛ㄦ寜 import_time 寤虹储寮曪紝鏀拺鏈€杩戝鍏ヨ拷韪煡璇€?,
            "瑙嗛琛?dedupe_key 鍞竴绱㈠紩锛屾帶鍒堕噸澶嶅啓鍏ユ垚鏈€?,
            "瑙嗛缁熻琛ㄥ敮涓€閿?video_id, stat_date)鍑忓皯閲嶅缁熻鏁版嵁銆?,
        ],
    )
    doc.add_paragraph(
        "鎬ц兘寤鸿锛氬綋鏁版嵁瑙勬ā杩涗竴姝ュ闀挎椂锛屽彲瀵?user_behavior銆乿ideo_statistics 鎸夋湀浠藉垎鍖猴紝"
        "骞跺皢楂橀鑱氬悎缁撴灉鍋氱墿鍖栨垨缂撳瓨銆?
    )

    doc.add_heading("6. 鏁版嵁涓€鑷存€т笌浜嬪姟璁捐", level=1)
    add_bullets(
        doc,
        [
            "瀵煎叆閫昏緫浣跨敤浜嬪姟锛岀‘淇?video/user/user_behavior 鍚屾鎴愬姛鎴栧洖婊氥€?,
            "瀵瑰箓绛夊啓鍏ヤ娇鐢?ON DUPLICATE KEY UPDATE锛岄檷浣庨噸澶嶈姹傚奖鍝嶃€?,
            "鏃у簱鍗囩骇鏃跺厛琛ュ瓧娈靛啀鍥炲～锛屾渶鍚庡缓绔嬬储寮曪紝閬垮厤涓棿鐘舵€佷笉涓€鑷淬€?,
        ],
    )

    doc.add_heading("7. 瀹夊叏涓庤繍缁村缓璁?, level=1)
    add_bullets(
        doc,
        [
            "鐢熶骇鐜涓嶅簲纭紪鐮佹暟鎹簱瀵嗙爜锛屽缓璁敼涓虹幆澧冨彉閲忔垨瀵嗛挜鏈嶅姟銆?,
            "鎵ц鏁版嵁搴撻噸寤鸿剼鏈墠蹇呴』鍋氬叏閲忓浠姐€?,
            "瀹氭湡瀵煎嚭 import_job 涓?source_trace 鎶ュ憡锛屽舰鎴愬璁℃潗鏂欍€?,
            "鏁忔劅鏁版嵁鑴辨晱鍚庡啀鐢ㄤ簬璁烘枃灞曠ず涓庡叕寮€绛旇京銆?,
        ],
    )

    doc.add_heading("8. 鍙墿灞曟柟鍚?, level=1)
    add_bullets(
        doc,
        [
            "鏂板鏍囩缁磋〃锛坱ag銆乿ideo_tag锛夊疄鐜板鏍囩鍒嗘瀽銆?,
            "鏂板浠诲姟璋冨害琛紙scheduler_job锛夌粺涓€绠＄悊瀹氭椂閲囬泦銆?,
            "鏂板瀹為獙琛紙ab_experiment锛夋敮鎸佹帹鑽愮瓥鐣ュ姣斻€?,
            "寮曞叆 OLAP 瀛樺偍鎴?ClickHouse 浠ユ敮鎸佹洿澶ц妯℃煡璇€?,
        ],
    )

    path = output_dir / "鏁版嵁搴撹璁¤鏄庝功.docx"
    doc.save(path)
    return path


def generate_doc3(output_dir: Path) -> Path:
    doc = Document()
    setup_styles(doc)

    add_title(doc, "瑙嗛缃戠珯鏁版嵁鍒嗘瀽骞冲彴绯荤粺姒傝璁捐璇存槑涔?)
    add_meta(
        doc,
        [
            "绯荤粺鍚嶇О锛歏ideo Analysis Platform",
            f"鏂囨。鏃ユ湡锛歿datetime.now().strftime('%Y-%m-%d')}",
            "鎶€鏈爤锛歏ue3 + ECharts + Three.js + Spring Boot + MyBatis + MySQL + Python",
        ],
    )

    doc.add_heading("1. 姒傝堪", level=1)
    doc.add_paragraph(
        "鏈郴缁熺敤浜庣粺涓€绠＄悊瑙嗛鏁版嵁閲囬泦銆佸叆搴撱€佸垎鏋愪笌鍙鍖栥€傜郴缁熻璁＄洰鏍囨槸锛氬湪鍗曟満閮ㄧ讲鏉′欢涓嬶紝"
        "鎻愪緵绋冲畾銆佸彲杩借釜銆佸彲鎵╁睍鐨勮棰戞暟鎹爺绌跺钩鍙般€?
    )
    add_bullets(
        doc,
        [
            "闈㈠悜瀵硅薄锛氭瘯涓氳璁＄爺绌躲€佽绋嬫紨绀恒€佹暟鎹垎鏋愬疄楠屻€?,
            "鏍稿績闂锛氬鏉ユ簮瑙嗛鏁版嵁濡備綍缁熶竴澶勭悊骞跺舰鎴愬彲瑙ｉ噴鍒嗘瀽缁撹銆?,
            "鎬讳綋鍘熷垯锛氭ā鍧楀寲銆佷綆鑰﹀悎銆佸彲鏇挎崲銆佸彲鑴氭湰鍖栭儴缃层€?,
        ],
    )

    doc.add_heading("2. 闇€姹傚垎鏋愭憳瑕?, level=1)
    doc.add_heading("2.1 鍔熻兘闇€姹?, level=2)
    add_bullets(
        doc,
        [
            "鏁版嵁閲囬泦锛氭敮鎸?URL 閲囬泦銆佹ā鎷熼噰闆嗐€佹枃鏈鍏ャ€佹枃浠跺鍏ャ€?,
            "鏁版嵁瀛樺偍锛氫繚瀛樿棰戙€佺敤鎴枫€佽涓恒€佺粺璁°€佸鍏ヤ换鍔°€佽仛绫荤粨鏋溿€?,
            "鏁版嵁鍒嗘瀽锛氱儹闂ㄦ帓琛屻€佸垎绫荤粺璁°€佷簰鍔ㄦ晥鐜囥€佸钩鍙版紡鏂椼€佺敤鎴风敾鍍忋€?,
            "鏁版嵁杩借釜锛氭樉绀烘潵婧怳RL銆佸鍏ユ柟寮忋€佽川閲忓垎銆佸鍏ユ椂闂翠笌浠诲姟淇℃伅銆?,
            "鍙鍖栵細鏀寔 ECharts 2D 鍥捐〃涓?Three.js 3D 瀵规瘮鍥俱€?,
            "閮ㄧ讲杩愮淮锛氭敮鎸佷竴閿垵濮嬪寲銆佽嚜娴嬨€佸叏鑷姩鍚姩涓庢姤鍛婂鍑恒€?,
        ],
    )
    doc.add_heading("2.2 闈炲姛鑳介渶姹?, level=2)
    add_bullets(
        doc,
        [
            "鍙敤鎬э細鏂扮幆澧冧笅鍙湪鐭椂闂村唴鍚姩銆?,
            "绋冲畾鎬э細鑴氭湰澶辫触鍙畾浣嶏紝鎺ュ彛瓒呮椂鏈夐噸璇曟満鍒躲€?,
            "涓€鑷存€э細閲嶅瀵煎叆涓嶅簲浜х敓缁熻鐣稿彉銆?,
            "鍙淮鎶ゆ€э細鍓嶅悗绔垎灞傛竻鏅帮紝SQL闆嗕腑绠＄悊銆?,
            "鍚堣鎬э細瀵归珮椋庨櫓鎶撳彇鎿嶄綔澧炲姞鎺堟潈纭鎻愮ず銆?,
        ],
    )

    doc.add_heading("3. 绯荤粺鎬讳綋鏋舵瀯璁捐", level=1)
    doc.add_paragraph("绯荤粺閲囩敤鈥滃墠绔睍绀哄眰 + 鍚庣鏈嶅姟灞?+ 鏁版嵁灞?+ 鏁版嵁澶勭悊灞傗€濈殑鍒嗗眰鏋舵瀯銆?)
    add_table(
        doc,
        ["灞傜骇", "涓昏鎶€鏈?, "鑱岃矗"],
        [
            ["灞曠ず灞?, "Vue3 + Vue Router + ECharts + Three.js", "椤甸潰浜や簰銆佸浘琛ㄦ覆鏌撱€侀噰闆嗗叆鍙ｃ€佹暟鎹拷韪睍绀?],
            ["鏈嶅姟灞?, "Spring Boot + MyBatis", "REST API銆佷笟鍔＄紪鎺掋€佺粺璁¤仛鍚堛€佸鍏ラ€昏緫"],
            ["鏁版嵁灞?, "MySQL 8.x", "涓氬姟鏁版嵁鎸佷箙鍖栥€佺储寮曚紭鍖栥€佺害鏉熶笌涓€鑷存€т繚闅?],
            ["澶勭悊灞?, "Python 鑴氭湰锛坰pider/analysis/scripts锛?, "閲囬泦瑙ｆ瀽銆佸垵濮嬪寲銆佽嚜娴嬨€佹姤鍛婂鍑?],
        ],
    )

    doc.add_heading("4. 妯″潡璁捐", level=1)
    doc.add_heading("4.1 鍓嶇妯″潡", level=2)
    add_bullets(
        doc,
        [
            "棣栭〉妯″潡锛氬钩鍙扮畝浠嬨€佹牳蹇冩寚鏍囥€佸叧閿礊瀵熴€佹潵婧愭瑙堛€?,
            "鐑棬瑙嗛妯″潡锛氭帓琛屽浘 + 鏄庣粏琛紝鍚潵婧愬钩鍙颁笌璐ㄩ噺鍒嗐€?,
            "鍒嗙被缁熻妯″潡锛氭敮鎸?2D/3D 鍒囨崲瑙傚療鍒嗙被鎾斁宸紓銆?,
            "浜掑姩鏁堢巼妯″潡锛氬垎鍖轰簰鍔ㄧ巼銆佸钩鍙版紡鏂椼€佸钩鍙版爣鍑嗗寲瀵规瘮銆?,
            "鐢ㄦ埛鍒嗘瀽妯″潡锛氳涓虹粨鏋勩€佹爣绛惧垎甯冦€佺敾鍍忓缓璁€佹槑缁嗚〃銆?,
            "鏁版嵁閲囬泦妯″潡锛歎RL/鏂囨湰/鏂囦欢瀵煎叆锛屾墽琛岀粨鏋滄棩蹇楋紝鏉ユ簮杩借釜琛ㄣ€?,
        ],
    )

    doc.add_heading("4.2 鍚庣妯″潡", level=2)
    add_bullets(
        doc,
        [
            "VideoController锛氭彁渚?`/video/*` 鍒嗘瀽鎺ュ彛锛坥verview/hot/category/user/funnel/benchmark/source-trace/insight锛夈€?,
            "CrawlerController锛氭彁渚?`/crawler/*` 閲囬泦涓庡鍏ユ帴鍙ｃ€?,
            "VideoService + VideoMapper锛氳仛鍚堟煡璇€佺敾鍍忚绠椼€佹礊瀵熺敓鎴愩€?,
            "CrawlerService锛氳В鏋愭枃鏈?鏂囦欢銆佸幓閲嶅悎骞躲€佸啓鍏ヤ换鍔℃棩蹇椼€?,
            "SchemaUpgradeRunner锛氬惎鍔ㄦ椂鑷姩琛ュ瓧娈点€佸缓琛ㄣ€佸洖濉€佸缓绱㈠紩銆?,
        ],
    )

    doc.add_heading("4.3 鏁版嵁澶勭悊涓庤剼鏈ā鍧?, level=2)
    add_bullets(
        doc,
        [
            "init_db_python.py锛氭棤 mysql CLI 鏉′欢涓嬪垵濮嬪寲鏁版嵁搴撱€?,
            "full_auto_launch.ps1锛氬叏鑷姩鍚姩骞舵墽琛岃嚜娴嬨€?,
            "one_click_no_cli.ps1锛氫竴閿垵濮嬪寲骞惰嚜娴嬨€?,
            "smoke_test.py锛氭暟鎹簱涓庢帴鍙ｅ彲杈炬€ч獙璇侊紝杈撳嚭鎶ュ憡銆?,
            "export_analysis_report.py锛氬鍑?CSV 涓?Markdown 鎶ュ憡銆?,
            "spider.py锛氶噰闆嗘ā鏉跨敓鎴愩€佹牎楠屻€佸叆搴撴祦绋嬨€?,
        ],
    )

    doc.add_heading("5. 鏍稿績涓氬姟娴佺▼璁捐", level=1)
    doc.add_heading("5.1 鏁版嵁閲囬泦鍏ュ簱娴佺▼", level=2)
    add_numbered(
        doc,
        [
            "鐢ㄦ埛鍦ㄥ墠绔噰闆嗛〉閫夋嫨瀵煎叆鏂瑰紡锛圲RL銆佹枃鏈€佹枃浠讹級銆?,
            "鍓嶇鎻愪氦璇锋眰鍒?CrawlerController銆?,
            "CrawlerService 瑙ｆ瀽鍘熷鍐呭锛岃瘑鍒钩鍙颁笌鍏抽敭瀛楁銆?,
            "鐢熸垚 dedupe_key 骞舵墽琛屽幓閲嶅悎骞讹紝璁＄畻璐ㄩ噺鍒嗐€?,
            "鍒涘缓 import_job锛屽啓鍏?video/user/user_behavior銆?,
            "杩斿洖浠诲姟缁撴灉锛屽墠绔睍绀烘棩蹇楀苟鍒锋柊鏉ユ簮杩借釜琛ㄣ€?,
        ],
    )
    doc.add_heading("5.2 鍒嗘瀽鏌ヨ娴佺▼", level=2)
    add_numbered(
        doc,
        [
            "鍓嶇椤甸潰鍔犺浇鏃惰皟鐢?`/video/*` 鎺ュ彛銆?,
            "Service 灞傝皟鐢?Mapper 鎵ц鑱氬悎 SQL銆?,
            "杩斿洖 VO 瀵硅薄缁欏墠绔浘琛ㄧ粍浠躲€?,
            "鐢ㄦ埛鍙湪鍚屼竴椤甸潰鍒囨崲 2D/3D 瑙嗗浘骞舵煡鐪嬭〃鏍兼槑缁嗐€?,
        ],
    )

    doc.add_heading("6. 鎺ュ彛姒傝璁捐", level=1)
    add_table(
        doc,
        ["鎺ュ彛", "鏂规硶", "璇存槑"],
        [
            ["/health", "GET", "绯荤粺鍋ュ悍妫€鏌?],
            ["/video/overview", "GET", "骞冲彴鎬昏鎸囨爣"],
            ["/video/hot", "GET", "鐑棬瑙嗛鎺掕"],
            ["/video/category", "GET", "鍒嗙被缁熻"],
            ["/video/engagement/category", "GET", "鍒嗙被浜掑姩鐜?],
            ["/video/engagement/video", "GET", "楂樹簰鍔ㄨ棰戝垪琛?],
            ["/video/user", "GET", "鐢ㄦ埛鐢诲儚鍒嗘瀽"],
            ["/video/platform", "GET", "骞冲彴瑙勬ā缁熻"],
            ["/video/funnel", "GET", "骞冲彴浜掑姩婕忔枟"],
            ["/video/platform/benchmark", "GET", "骞冲彴鏍囧噯鍖栧姣?],
            ["/video/source-trace", "GET", "鏉ユ簮杩借釜璁板綍"],
            ["/video/insight", "GET", "鑷姩娲炲療鍗＄墖"],
            ["/crawler/run-url", "POST", "URL閲囬泦浠诲姟"],
            ["/crawler/run-mock", "POST", "妯℃嫙閲囬泦浠诲姟"],
            ["/crawler/import-text", "POST", "鏂囨湰鏅鸿兘瀵煎叆"],
            ["/crawler/import-file", "POST", "鏂囦欢鏅鸿兘瀵煎叆"],
        ],
    )

    doc.add_heading("7. 閮ㄧ讲姒傝璁捐", level=1)
    add_bullets(
        doc,
        [
            "鍚庣鏈嶅姟绔彛锛?080锛涘墠绔紑鍙戠鍙ｏ細5173锛涙暟鎹簱绔彛锛?306銆?,
            "榛樿鏁版嵁搴擄細video_analysis锛坮oot/你的MySQL密码锛夈€?,
            "鎺ㄨ崘鍚姩鏂瑰紡锛氭牴鐩綍鍙屽嚮涓€閿剼鏈紝鑷姩瀹屾垚鍒濆鍖栦笌楠岃瘉銆?,
            "鎶ュ憡杈撳嚭鐩綍锛歚analysis/reports`锛屽惈 latest_smoke 涓?latest_*_report 鏂囦欢銆?,
        ],
    )

    doc.add_heading("8. 瀹夊叏涓庨闄╂帶鍒惰璁?, level=1)
    add_bullets(
        doc,
        [
            "楂橀闄?URL 鎶撳彇榛樿闇€瑕侀闄╃‘璁わ紝涓嶉紦鍔辨湭鎺堟潈鏁版嵁鎶撳彇銆?,
            "瀵煎叆鏃ュ織鍜屾潵婧愯拷韪敤浜庡璁★紝鏀寔瀹氫綅寮傚父鏁版嵁鏉ユ簮銆?,
            "鑴氭湰绾ф姤閿欎俊鎭槑纭紝渚夸簬蹇€熸帓闅滐紙鐜銆佷緷璧栥€佽繛鎺ャ€佽秴鏃讹級銆?,
            "鏁版嵁搴撻噸寤鸿剼鏈渶鏄庣‘鎻愮ず浼氭竻绌哄師搴擄紝闃叉璇搷浣溿€?,
        ],
    )

    doc.add_heading("9. 娴嬭瘯涓庨獙鏀舵瑕?, level=1)
    add_table(
        doc,
        ["娴嬭瘯绫诲瀷", "鐩爣", "涓昏鏂瑰紡", "楠屾敹鏍囧噯"],
        [
            ["鏋勫缓娴嬭瘯", "淇濊瘉浠ｇ爜鍙紪璇?, "mvn compile / npm run build", "鏃犵紪璇戦敊璇?],
            ["鎺ュ彛娴嬭瘯", "淇濊瘉鍚庣鍙敤", "smoke_test.py", "鍏抽敭鎺ュ彛杩斿洖鎴愬姛鎴栧彲瑙ｉ噴鍛婅"],
            ["鍔熻兘娴嬭瘯", "淇濊瘉涓氬姟瀹屾暣", "椤甸潰鍏ㄦ祦绋嬫墜宸ユ祴璇?, "閲囬泦->鍏ュ簱->鍒嗘瀽->杩借釜閾捐矾閫氱晠"],
            ["鏁版嵁娴嬭瘯", "淇濊瘉缁熻鍙俊", "鍘婚噸銆佽川閲忓垎銆佸鍏ヤ换鍔℃牳鏌?, "鏃犳槑鏄鹃噸澶嶈啫鑳€涓庢潵婧愮己澶?],
        ],
    )

    doc.add_heading("10. 鍚庣画婕旇繘寤鸿", level=1)
    add_bullets(
        doc,
        [
            "寮曞叆閴存潈涓庤鑹叉潈闄愶紙绠＄悊鍛?鍒嗘瀽甯?璁垮锛夈€?,
            "寮曞叆浠诲姟璋冨害涓績锛屽疄鐜板畾鏃堕噰闆嗗拰澶辫触閲嶈瘯銆?,
            "琛ュ厖鏁版嵁琛€缂樺彲瑙嗗寲锛屾彁鍗囩爺绌跺彲瑙ｉ噴鎬с€?,
            "澧炲姞娴嬭瘯瑕嗙洊鐜囷紙鍗曞厓娴嬭瘯銆佹帴鍙ｅ绾︽祴璇曘€丒2E锛夈€?,
            "浼樺寲鍓嶇鏋勫缓浣撶Н锛屾寜璺敱鎷嗗寘鎻愬崌鍔犺浇鎬ц兘銆?,
        ],
    )

    path = output_dir / "绯荤粺姒傝璁捐璇存槑涔?docx"
    doc.save(path)
    return path


def main() -> None:
    output_dir = Path(r"C:\Users\17929\Desktop\姣曚笟璁捐")
    output_dir.mkdir(parents=True, exist_ok=True)

    p1 = generate_doc1(output_dir)
    p2 = generate_doc2(output_dir)
    p3 = generate_doc3(output_dir)

    print(str(p1))
    print(str(p2))
    print(str(p3))


if __name__ == "__main__":
    main()


